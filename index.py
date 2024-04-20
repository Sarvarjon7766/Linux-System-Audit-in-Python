import subprocess
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog
from docx import Document

def update_data():
    # Barcha ishlayotgan dasturlarni olish
    output = subprocess.check_output(['ps', '-eo', 'user,pid,comm,%cpu'])

    lines = output.decode().split('\n')

    # Yangilash uchun jadvalni tozalash
    for i in tree.get_children():
        tree.delete(i)

    # Jadvalga ma'lumotlarni qo'shish
    for line in lines[1:]:
        parts = line.split()
        if len(parts) >= 4 and float(parts[-1]) > 0.0:
            user = parts[0]
            pid = parts[1]
            command = " ".join(parts[2:-1])
            cpu = parts[-1]
            tree.insert("", "end", values=(user, pid, command, cpu))

    # Toshkent vaqti bilan har 1 sekundda yangilash
    if running.get():
        root.after(1000, update_data)

def start_update():
    running.set(True)
    update_data()

def stop_update():
    running.set(False)

def sort_data():
    pid_or_name = search_entry.get()
    if pid_or_name:
        # Filter data by PID or process name
        filtered_data = [line.split() for line in subprocess.check_output(['ps', '-eo', 'user,pid,comm,%cpu']).decode().split('\n')[1:] if pid_or_name in line]
        if filtered_data:
            # Create a new window to display sorted data
            sort_window = tk.Toplevel(root)
            sort_window.title("Sorted Data")
            sorted_tree = ttk.Treeview(sort_window, columns=("Foydalanuvchi", "PID", "Dastur nomi", "CPU(%)"), show="headings")
            sorted_tree.heading("Foydalanuvchi", text="Foydalanuvchi")
            sorted_tree.heading("PID", text="PID")
            sorted_tree.heading("Dastur nomi", text="Dastur nomi")
            sorted_tree.heading("CPU(%)", text="CPU(%)")
            sorted_tree.pack(expand=True, fill=tk.BOTH)
            for data in filtered_data:
                user, pid, command, cpu = data[0], data[1], " ".join(data[2:-1]), data[-1]
                sorted_tree.insert("", "end", values=(user, pid, command, cpu))
        else:
            messagebox.showinfo("Ma'lumot topilmadi", "Kiritilgan PID yoki dastur nomi bo'yicha ma'lumot topilmadi.")
    else:
        messagebox.showinfo("Ma'lumot kiritilmadi", "Iltimos, qidirish uchun PID yoki dastur nomini kiriting.")

def save_data():
    # Qayerga saqlash kerakligini so'ramiz
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc = Document()
        doc.add_heading("Ishlayotgan Dasturlar", level=1)
        for item in tree.get_children():
            values = tree.item(item, 'values')
            doc.add_paragraph(f"Foydalanuvchi: {values[0]}, PID: {values[1]}, Dastur nomi: {values[2]}, CPU(%): {values[3]}")
        doc.save(file_path)

# Asosiy oynani yaratish
root = tk.Tk()
root.title("Ishlayotgan Dasturlar")

running = tk.BooleanVar()
search_entry = tk.Entry(root, width=30)
search_entry.pack(pady=10)

button_frame = tk.Frame(root)
button_frame.pack(pady=5)

start_button = tk.Button(button_frame, text="Boshlash", command=start_update)
start_button.pack(side=tk.LEFT, padx=5, pady=5)

stop_button = tk.Button(button_frame, text="To'xtatish", command=stop_update)
stop_button.pack(side=tk.LEFT, padx=5, pady=5)

sort_button = tk.Button(button_frame, text="Saralash", command=sort_data)
sort_button.pack(side=tk.LEFT, padx=5, pady=5)

save_button = tk.Button(button_frame, text="Saqlash", command=save_data)
save_button.pack(side=tk.LEFT, padx=5, pady=5)

# Jadvalni yaratish
tree = ttk.Treeview(root, columns=("Foydalanuvchi", "PID", "Dastur nomi", "CPU(%)"), show="headings")
tree.heading("Foydalanuvchi", text="Foydalanuvchi")
tree.heading("PID", text="PID")
tree.heading("Dastur nomi", text="Dastur nomi")
tree.heading("CPU(%)", text="CPU(%)")
tree.pack(expand=True, fill=tk.BOTH)

# Ma'lumotlarni yangilash
update_data()

# Dasturni ishga tushirish
root.mainloop()
